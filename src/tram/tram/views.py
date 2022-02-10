import io
import json
import re

from constance import config
from django.contrib.auth.decorators import login_required
from django.http import (
    Http404,
    HttpResponse,
    HttpResponseBadRequest,
    StreamingHttpResponse,
)
from django.shortcuts import render
from django.utils.text import slugify
from rest_framework import viewsets

from tram import serializers
from tram.ml import base
from tram.models import AttackObject, DocumentProcessingJob, Mapping, Report, Sentence

from docx import Document
from docx.shared import Inches


class AttackObjectViewSet(viewsets.ModelViewSet):
    queryset = AttackObject.objects.all()
    serializer_class = serializers.AttackObjectSerializer


class DocumentProcessingJobViewSet(viewsets.ModelViewSet):
    queryset = DocumentProcessingJob.objects.all()
    serializer_class = serializers.DocumentProcessingJobSerializer


class MappingViewSet(viewsets.ModelViewSet):
    queryset = Mapping.objects.all()
    serializer_class = serializers.MappingSerializer

    def get_queryset(self):
        queryset = MappingViewSet.queryset
        sentence_id = self.request.query_params.get("sentence-id", None)
        if sentence_id:
            queryset = queryset.filter(sentence__id=sentence_id)

        return queryset


class ReportViewSet(viewsets.ModelViewSet):
    queryset = Report.objects.all()
    serializer_class = serializers.ReportSerializer


class ReportExportViewSet(viewsets.ModelViewSet):
    queryset = Report.objects.all()
    serializer_class = serializers.ReportExportSerializer

    def retrieve(self, request, *args, **kwargs):

        format = request.GET.get("type", "")

        # If an invalid format is given, just default to json
        if format not in ["json", "docx"]:
            format = "json"
            print("Invalid File Type. Defaulting to JSON.")

        # Retrieve report data as json
        response = super().retrieve(request, *args, **kwargs)

        if format == "json":
            filename = slugify(self.get_object().name) + ".json"
            response["Content-Disposition"] = 'attachment; filename="%s"' % filename

        elif format == "docx":
            # Uses json dictionary to create formatted document
            document = self.build_docx(response.data)

            # save document info
            buffer = io.BytesIO()
            document.save(buffer)  # save your memory stream
            buffer.seek(0)  # rewind the stream

            # put them to streaming content response within docx content_type
            content_type = (
                "application/"
                "vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            response = StreamingHttpResponse(
                streaming_content=buffer,  # use the stream's content
                content_type=content_type,
            )

            filename = slugify(self.get_object().name) + ".docx"
            response["Content-Disposition"] = 'attachment; filename="%s"' % filename
            response["Content-Encoding"] = "UTF-8"

        return response

    # Uses json dictionary of the report to build a formatted document
    def build_docx(self, data):
        document = Document()
        name = data["name"]
        accepted = str(data["accepted_sentences"])
        reviewing = str(data["reviewing_sentences"])
        total = str(data["total_sentences"])
        text = data["text"]
        sentences = data["sentences"]
        accepted_sentences = [s for s in sentences if s["mappings"]]

        # Display header with basic stats
        document.add_heading("TRAM " + name)
        paragraph = document.add_paragraph("")
        paragraph.add_run("Accepted Sentences: ").bold = True
        paragraph.add_run(accepted + "\n")
        paragraph.add_run("Reviewing Sentences: ").bold = True
        paragraph.add_run(reviewing + "\n")
        paragraph.add_run("Total Sentences: ").bold = True
        paragraph.add_run(total)

        # Display all attack techniques found in the text
        document.add_heading("Techniques Found")
        techniques = set()

        # Find all attack techniques in the report
        for sentence in accepted_sentences:
            for mapping in sentence["mappings"]:
                curMapping = (mapping["attack_id"], mapping["name"])
                if not curMapping in techniques:
                    techniques.add(curMapping)

        # Sort attack techniques by integer part
        techniques = sorted(techniques, key=lambda x: float(x[0][1:]))
        num_techniques = len(techniques)

        # Display all attack techniques
        paragraph = document.add_paragraph("")
        paragraph.add_run("Total Techniques: ").bold = True
        paragraph.add_run(str(num_techniques) + "\n")
        for technique in techniques:
            paragraph.add_run("Attack Id: ").bold = True
            paragraph.add_run(technique[0])
            paragraph.add_run(", Name: ").bold = True
            paragraph.add_run(technique[1] + "\n")

        # Display matched sentences in a table
        document.add_page_break()
        document.add_heading("Matched Sentences", level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = "TableGrid"
        table.autofit = False
        table.allow_autofit = False

        # This resizing format is strange, works for now
        table.columns[0].width = Inches(4.0)
        table.columns[1].width = Inches(3.0)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Text"
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].text = "Mappings"
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True

        for sentence in accepted_sentences:
            row_cells = table.add_row().cells
            row_cells[0].text = re.sub(r"[\n\r]*", "", sentence["text"])
            row_cells[1].text = ""
            map_paragraph = row_cells[1].paragraphs[0]

            for mapping in sentence["mappings"]:
                if sentence["disposition"] == "accept":
                    confidence = "Manually Accepted"
                else:
                    confidence = mapping["confidence"] + "%"

                map_paragraph.add_run("Attack Id: ").bold = True
                map_paragraph.add_run(mapping["attack_id"] + ", ")
                map_paragraph.add_run("Name: ").bold = True
                map_paragraph.add_run(mapping["name"] + ", ")
                map_paragraph.add_run("Confidence: ").bold = True
                map_paragraph.add_run(confidence + "\n")

        # Display full text, remove extra white space to increase readability
        document.add_page_break()
        document.add_heading("Full Document", level=1)
        document.add_paragraph(re.sub(r"[\r]*", "", text))

        return document


class SentenceViewSet(viewsets.ModelViewSet):
    queryset = Sentence.objects.all()
    serializer_class = serializers.SentenceSerializer

    def get_queryset(self):
        queryset = SentenceViewSet.queryset
        report_id = self.request.query_params.get("report-id", None)
        if report_id:
            queryset = queryset.filter(report__id=report_id)

        attack_id = self.request.query_params.get("attack-id", None)
        if attack_id:
            sentences = Mapping.objects.filter(
                attack_object__attack_id=attack_id
            ).values("sentence")
            queryset = queryset.filter(id__in=sentences)

        return queryset


@login_required
def index(request):
    jobs = DocumentProcessingJob.objects.all()
    job_serializer = serializers.DocumentProcessingJobSerializer(jobs, many=True)

    reports = Report.objects.all()
    report_serializer = serializers.ReportSerializer(reports, many=True)

    context = {
        "job_queue": job_serializer.data,
        "reports": report_serializer.data,
    }

    return render(request, "index.html", context=context)


@login_required
def upload(request):
    """Places a file into ml-pipeline for analysis"""
    if request.method != "POST":
        return HttpResponse("Request method must be POST", status=405)

    file_content_type = request.FILES["file"].content_type
    if file_content_type in (
        "application/pdf",  # .pdf files
        "text/html",  # .html files
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx files
    ):
        DocumentProcessingJob.create_from_file(request.FILES["file"])
    elif file_content_type in ("application/json",):  # .json files
        json_data = json.loads(request.FILES["file"].read())
        res = serializers.ReportExportSerializer(data=json_data)

        if res.is_valid():
            res.save()
        else:
            return HttpResponseBadRequest(res.errors)
    else:
        return HttpResponseBadRequest("Unsupported file type")

    return HttpResponse("File saved for processing", status=200)


@login_required
def ml_home(request):
    techniques = AttackObject.get_sentence_counts()
    model_metadata = base.ModelManager.get_all_model_metadata()

    context = {
        "techniques": techniques,
        "ML_ACCEPT_THRESHOLD": config.ML_ACCEPT_THRESHOLD,
        "ML_CONFIDENCE_THRESHOLD": config.ML_CONFIDENCE_THRESHOLD,
        "models": model_metadata,
    }

    return render(request, "ml_home.html", context)


@login_required
def ml_technique_sentences(request, attack_id):
    context = {"attack_id": attack_id}
    return render(request, "technique_sentences.html", context)


@login_required
def ml_model_detail(request, model_key):
    try:
        model_metadata = base.ModelManager.get_model_metadata(model_key)
    except ValueError:
        raise Http404("Model does not exists")
    context = {"model": model_metadata}
    return render(request, "model_detail.html", context)


@login_required
def analyze(request, pk):
    report = Report.objects.get(id=pk)
    techniques = AttackObject.objects.all().order_by("attack_id")
    tecniques_serializer = serializers.AttackObjectSerializer(techniques, many=True)

    context = {
        "report_id": report.id,
        "report_name": report.name,
        "attack_techniques": tecniques_serializer.data,
    }
    return render(request, "analyze.html", context)
