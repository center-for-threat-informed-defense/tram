import re

from docx import Document
from docx.shared import Inches


def build(data):
    """Convert JSON report into Word document (.docx)."""
    document = Document()
    name = data["name"]
    accepted = str(data["accepted_sentences"])
    reviewing = str(data["reviewing_sentences"])
    total = str(data["total_sentences"])
    text = data["text"]
    sentences = data["sentences"]
    accepted_sentences = [s for s in sentences if s["mappings"]]

    # Display header with basic stats
    document.add_heading("TRAM " + name)
    paragraph = document.add_paragraph("")
    paragraph.add_run("Accepted Sentences: ").bold = True
    paragraph.add_run(accepted + "\n")
    paragraph.add_run("Reviewing Sentences: ").bold = True
    paragraph.add_run(reviewing + "\n")
    paragraph.add_run("Total Sentences: ").bold = True
    paragraph.add_run(total)

    # Display all attack techniques found in the text
    document.add_heading("Techniques Found")
    techniques = set()

    # Find all attack techniques in the report
    for sentence in accepted_sentences:
        for mapping in sentence["mappings"]:
            curMapping = (mapping["attack_id"], mapping["name"])
            if curMapping not in techniques:
                techniques.add(curMapping)

    # Sort attack techniques by integer part
    techniques = sorted(techniques, key=lambda x: float(x[0][1:]))
    num_techniques = len(techniques)

    # Display all attack techniques
    paragraph = document.add_paragraph("")
    paragraph.add_run("Total Techniques: ").bold = True
    paragraph.add_run(str(num_techniques) + "\n")
    for technique in techniques:
        paragraph.add_run("Attack Id: ").bold = True
        paragraph.add_run(technique[0])
        paragraph.add_run(", Name: ").bold = True
        paragraph.add_run(technique[1] + "\n")

    # Display matched sentences in a table
    document.add_page_break()
    document.add_heading("Matched Sentences", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "TableGrid"
    table.autofit = False
    table.allow_autofit = False

    # This resizing format is strange, works for now
    table.columns[0].width = Inches(4.0)
    table.columns[1].width = Inches(3.0)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Text"
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].text = "Mappings"
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True

    for sentence in accepted_sentences:
        row_cells = table.add_row().cells
        row_cells[0].text = re.sub(r"[\n\r]*", "", sentence["text"])
        row_cells[1].text = ""
        map_paragraph = row_cells[1].paragraphs[0]

        for mapping in sentence["mappings"]:
            if sentence["disposition"] == "accept":
                confidence = "Manually Accepted"
            else:
                confidence = mapping["confidence"] + "%"

            map_paragraph.add_run("Attack Id: ").bold = True
            map_paragraph.add_run(mapping["attack_id"] + ", ")
            map_paragraph.add_run("Name: ").bold = True
            map_paragraph.add_run(mapping["name"] + ", ")
            map_paragraph.add_run("Confidence: ").bold = True
            map_paragraph.add_run(confidence + "\n")

    # Add the full text of the report to the finished document. Remove control
    # characters, otherwise the docx library will raise an error.
    document.add_page_break()
    document.add_heading("Full Document", level=1)
    cleaned_text = re.sub("[\x00-\x09\x0B-\x1F\x7F]", "�", text)
    document.add_paragraph(cleaned_text)

    return document
